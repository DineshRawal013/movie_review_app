"""
Markdown → .docx batch converter
Converts all specified MD files to professional Word documents with:
  - Title page
  - Auto Table of Contents placeholder
  - Consistent Heading 1/2/3 styles
  - Header: project branding
  - Footer: document name, version, page number
  - Tables with alternating row shading
  - Code blocks as styled paragraphs
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Paths ────────────────────────────────────────────────────────────────────
DOCS     = Path(__file__).parent.parent
OUT_DIR  = Path(__file__).parent

PROJECT  = "Movie Review App"
COMPANY  = "AI Dev Team"
VERSION  = "v1.0.0"
DATE     = "2026-05-26"

# ── Colour palette ───────────────────────────────────────────────────────────
NAVY      = RGBColor(0x1F, 0x38, 0x64)
BLUE      = RGBColor(0x2E, 0x75, 0xB6)
LIGHT_BLUE= RGBColor(0xD6, 0xE4, 0xF0)
TABLE_HDR = RGBColor(0x1F, 0x38, 0x64)
ROW_ALT   = RGBColor(0xF2, 0xF7, 0xFF)
CODE_BG   = RGBColor(0xF5, 0xF5, 0xF5)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)


# ── Helpers: low-level XML ────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def add_page_number(run):
    """Insert PAGE field into run."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_toc(doc):
    """Insert a TOC field that Word will populate on open."""
    para = doc.add_paragraph()
    para.style = doc.styles['TOC Heading'] if 'TOC Heading' in doc.styles else para.style
    run = para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = r'TOC \o "1-3" \h \z \u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return para


def set_header_footer(doc, doc_title: str, version: str):
    section = doc.sections[0]
    section.different_first_page_header_footer = False

    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run(f"{PROJECT}  |  {doc_title}")
    run.font.name  = "Calibri"
    run.font.size  = Pt(9)
    run.font.color.rgb = BLUE
    run.font.bold  = True
    # bottom border on header paragraph
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E75B6')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run(f"{COMPANY}  |  {version}  |  {DATE}  |  Page ").font.size = Pt(9)
    pn_run = fp.add_run()
    pn_run.font.size = Pt(9)
    add_page_number(pn_run)


def add_title_page(doc, doc_title: str, subtitle: str, version: str):
    """Insert a professional title page."""
    # Big coloured block via paragraph shading
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  '1F3864')
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(120)
    p.paragraph_format.space_after  = Pt(0)

    # Project name
    r = p.add_run(PROJECT)
    r.font.name   = "Calibri"
    r.font.size   = Pt(28)
    r.font.bold   = True
    r.font.color.rgb = WHITE

    # Document title
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr2 = p2._p.get_or_add_pPr()
    shd2 = OxmlElement('w:shd')
    shd2.set(qn('w:val'), 'clear'); shd2.set(qn('w:color'), 'auto'); shd2.set(qn('w:fill'), '1F3864')
    pPr2.append(shd2)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(48)
    r2 = p2.add_run(doc_title)
    r2.font.name  = "Calibri"
    r2.font.size  = Pt(20)
    r2.font.color.rgb = LIGHT_BLUE

    # Metadata table
    meta = doc.add_table(rows=4, cols=2)
    meta.style = 'Table Grid'
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, (k, v) in enumerate([
        ("Version", version),
        ("Date",    DATE),
        ("Project", PROJECT),
        ("Prepared by", COMPANY),
    ]):
        meta.rows[row].cells[0].text = k
        meta.rows[row].cells[1].text = v
        meta.rows[row].cells[0].paragraphs[0].runs[0].bold = True
        set_cell_bg(meta.rows[row].cells[0], 'D6E4F0')

    doc.add_page_break()


# ── Style setup ──────────────────────────────────────────────────────────────
def setup_styles(doc):
    styles = doc.styles

    def ensure(name, base_name='Normal'):
        if name not in [s.name for s in styles]:
            s = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            s.base_style = styles[base_name]
            return s
        return styles[name]

    # Normal
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    # Heading 1
    h1 = styles['Heading 1']
    h1.font.name  = 'Calibri'
    h1.font.size  = Pt(18)
    h1.font.bold  = True
    h1.font.color.rgb = NAVY
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after  = Pt(6)

    # Heading 2
    h2 = styles['Heading 2']
    h2.font.name  = 'Calibri'
    h2.font.size  = Pt(14)
    h2.font.bold  = True
    h2.font.color.rgb = BLUE
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after  = Pt(4)

    # Heading 3
    h3 = styles['Heading 3']
    h3.font.name  = 'Calibri'
    h3.font.size  = Pt(12)
    h3.font.bold  = True
    h3.font.color.rgb = BLUE
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after  = Pt(2)

    # Code style
    code = ensure('Code Block', 'Normal')
    code.font.name = 'Courier New'
    code.font.size = Pt(9)
    code.paragraph_format.left_indent  = Cm(1)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after  = Pt(4)


# ── Inline markdown → runs ───────────────────────────────────────────────────
def add_inline(para, text: str):
    """Parse **bold**, *italic*, `code` and plain text into runs."""
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            para.add_run(text[last:m.start()])
        full = m.group(0)
        if full.startswith('**'):
            r = para.add_run(m.group(2))
            r.bold = True
        elif full.startswith('*'):
            r = para.add_run(m.group(3))
            r.italic = True
        else:  # backtick code
            r = para.add_run(m.group(4))
            r.font.name = 'Courier New'
            r.font.size = Pt(10)
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])


# ── Markdown → Word ───────────────────────────────────────────────────────────
def md_to_docx(md_text: str, doc: Document):
    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_lines = []
    skip_meta = True  # skip front-matter key:value lines at very top

    while i < len(lines):
        line = lines[i]

        # Skip the document title line (we have title page)
        if i == 0 and line.startswith('# '):
            i += 1
            continue

        # Skip metadata key:value lines at top
        if skip_meta and re.match(r'^\*\*\w[\w\s]+\*\*:', line):
            i += 1
            continue
        else:
            skip_meta = False

        # Code fence
        if line.strip().startswith('```'):
            if in_code:
                # End code block
                for cl in code_lines:
                    p = doc.add_paragraph(style='Code Block')
                    p.add_run(cl)
                    pPr = p._p.get_or_add_pPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'F5F5F5')
                    pPr.append(shd)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Headings
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            text  = m.group(2).strip()
            # strip trailing {#anchor}
            text = re.sub(r'\s*\{#[^}]+\}', '', text)
            style_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3',
                         4: 'Heading 3', 5: 'Heading 3', 6: 'Heading 3'}
            p = doc.add_paragraph(style=style_map[level])
            add_inline(p, text)
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^[-*_]{3,}\s*$', line):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bot  = OxmlElement('w:bottom')
            bot.set(qn('w:val'), 'single')
            bot.set(qn('w:sz'), '6')
            bot.set(qn('w:space'), '1')
            bot.set(qn('w:color'), '2E75B6')
            pBdr.append(bot)
            pPr.append(pBdr)
            i += 1
            continue

        # Table — look-ahead for separator
        if '|' in line and i + 1 < len(lines) and re.match(r'\s*\|[-| :]+\|\s*$', lines[i+1]):
            headers = [c.strip() for c in line.strip().strip('|').split('|')]
            i += 2  # skip separator
            table_rows = [headers]
            while i < len(lines) and '|' in lines[i]:
                cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                while len(cells) < len(headers):
                    cells.append('')
                table_rows.append(cells[:len(headers)])
                i += 1
            # Build Word table
            t = doc.add_table(rows=len(table_rows), cols=len(headers))
            t.style = 'Table Grid'
            t.alignment = WD_TABLE_ALIGNMENT.LEFT
            for r_idx, row_data in enumerate(table_rows):
                for c_idx, cell_text in enumerate(row_data):
                    cell = t.rows[r_idx].cells[c_idx]
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                    p = cell.paragraphs[0]
                    p.clear()
                    if r_idx == 0:
                        add_inline(p, cell_text)
                        p.runs[0].bold = True if p.runs else None
                        p.runs[0].font.color.rgb = WHITE if p.runs else None
                        set_cell_bg(cell, '1F3864')
                    else:
                        add_inline(p, cell_text)
                        if r_idx % 2 == 0:
                            set_cell_bg(cell, 'F2F7FF')
            doc.add_paragraph()
            continue

        # Blockquote
        if line.startswith('>'):
            text = re.sub(r'^>\s?', '', line)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            add_inline(p, text)
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        # Unordered list
        m_ul = re.match(r'^(\s*)[-*+]\s+(.*)', line)
        if m_ul:
            indent = len(m_ul.group(1)) // 2
            text   = m_ul.group(2)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
            add_inline(p, text)
            i += 1
            continue

        # Ordered list
        m_ol = re.match(r'^(\s*)\d+\.\s+(.*)', line)
        if m_ol:
            text = m_ol.group(2)
            p = doc.add_paragraph(style='List Number')
            add_inline(p, text)
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        add_inline(p, line)
        i += 1


# ── Per-document conversion ───────────────────────────────────────────────────
DOCS_TO_CONVERT = [
    ("SRS.md",           "SRS.docx",             "Software Requirements Specification", "v1.0"),
    ("HLD.md",           "HLD.docx",              "High-Level Design",                   "v1.0"),
    ("LLD.md",           "LLD.docx",              "Low-Level Design",                    "v1.0"),
    ("TEST-PLAN.md",     "TestPlan.docx",          "Master Test Plan",                    "v1.0"),
    ("USER-GUIDE.md",    "UserGuide.docx",         "User Guide",                          "v1.0"),
    ("API-REFERENCE.md", "API-Reference.docx",     "API Reference",                       "v1.0"),
    ("CLOSURE.md",       "ProjectClosure.docx",    "Project Closure Report",              "v1.0"),
]


def convert_one(md_file: Path, out_file: Path, doc_title: str, version: str):
    print(f"  Converting {md_file.name} → {out_file.name} ...", end=" ", flush=True)
    text = md_file.read_text(encoding="utf-8")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.54)
        section.right_margin  = Cm(2.54)

    setup_styles(doc)
    add_title_page(doc, doc_title, "", version)
    set_header_footer(doc, doc_title, version)

    # TOC page
    toc_heading = doc.add_paragraph("Table of Contents", style='Heading 1')
    add_toc(doc)
    doc.add_page_break()

    md_to_docx(text, doc)

    doc.save(out_file)
    print("✅")


def main():
    print(f"\n{'='*60}")
    print(f"  Movie Review App — Document Conversion")
    print(f"  Output → {OUT_DIR}")
    print(f"{'='*60}\n")

    for md_name, docx_name, title, ver in DOCS_TO_CONVERT:
        md_path   = DOCS / md_name
        docx_path = OUT_DIR / docx_name
        if not md_path.exists():
            print(f"  ⚠️  SKIP {md_name} (not found)")
            continue
        try:
            convert_one(md_path, docx_path, title, ver)
        except Exception as e:
            print(f"  ❌  ERROR: {e}")

    print(f"\n✅  All done. Files in {OUT_DIR}\n")


if __name__ == "__main__":
    main()
